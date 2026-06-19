#!/usr/bin/env python3
"""
Change Order Generator — v3
Grounded in CHG-51 (Change Event SOP) and SYS-09 (Estimating Schedule).

v3 changes:
- Claude generates line-item cost breakdown using SYS-09 pricing database
- Contractor's labor/material inputs serve as anchors, not the only numbers
- Cost section now shows item-by-item breakdown with cost basis references
- Markup applied to Claude-generated subtotal
"""

import os
from datetime import datetime, timedelta
from pathlib import Path

import anthropic
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MARKUP = {
    "application_fee_pct": 0.01,
    "overhead_pct":        0.10,
    "fee_pct":             0.10,
    "total_pct":           0.21,
}

PCO_VALIDITY_DAYS = 21

CAUSE_CATEGORIES = [
    "Owner-Requested Change",
    "Design Revision / ASI",
    "Unforeseen Field Condition",
    "Allowance Over/Under",
    "Coordination Conflict",
    "Unit Price Work",
    "Other",
]


# ---------------------------------------------------------------------------
# SYS-09 Pricing Reference (embedded for Claude context)
# ---------------------------------------------------------------------------

SYS09_RATES = """
FIRM PRICING DATABASE — SYS-09 (Beacon Hill Interiors)
Based on 12+ completed projects, 1,258 consolidated line items.
Use these rates to build line-item cost breakdowns. Reference as "SYS-09" in cost basis column.

LABOR RATES:
- Demo crew: $1,920/day (any demolition task)
- Carpentry crew (rough or finish): $2,486/day
- General laborer: $75–85/hr
- Skilled tradesperson: $89/hr

MARKUP STRUCTURE (apply after construction subtotal):
- Application Fee: 1%
- Overhead: 10%
- Fee: 10%
- Total markup: 21%

DIV 1 — GENERAL REQUIREMENTS:
- Trash removal, single axle dump truck: $1,301/load
- Site protection (floors + dust barriers): $1,000–3,000 LS
- Daily cleaning: $75–85/hr
- Demo materials package: $200–500 LS

DIV 2 — DEMOLITION:
- Any demolition task (tile, drywall, flooring, cabinetry, fixtures): $1,920/day (demo crew)
- Salvage and protect existing door/item for reuse: $200/EA

DIV 5 — METALS:
- Metal fasteners, joist hangers, bolts, screws package: $720–765 allowance

DIV 6 — CARPENTRY & MILLWORK:
- 3/4" subfloor material: $3.26–3.76/SF
- 3/4" subfloor install labor: $5.25–6.25/SF
- Blocking material: $3.26–3.76/SF; install: $5.25–6.25/SF
- Rough/finish carpentry crew (framing, millwork, closets, cabinets): $2,486/day
- Baseboard material (mid-grade): $6.00–7.25/LF
- Baseboard install: carpentry crew day rate
- Door casing material: $275–385/EA; install: $375–425/EA
- Window casing material: $350–425/EA; install: $375–450/EA
- Crown moulding material: $8–12/LF
- Cedar T&G planking material: $5–8/SF
- 2x4 interior wall framing: $38.76/LF
- 2x6 exterior wall framing: $48.92/LF
- Kitchen cabinetry (semi-custom): $10,000–45,000 (varies widely)
- Closet system (standard): $1,500–3,000

DIV 7 — THERMAL & MOISTURE:
- Sound insulation material: $2.50/SF; install: $1.50/SF
- Spray foam material: $2.00–2.25/SF; install: $1.25/SF

DIV 8 — DOORS & WINDOWS:
- Interior door (standard): $800–1,200 furnished
- Door hardware: $165–225/EA
- Install interior door: carpentry crew, typically 2–4 hrs

DIV 9 — FINISHES:
- Tile floor: $10/SF material + $28/SF install (total ~$38–45/SF)
- Hardwood new install: $8–20/SF material + $7/SF install + $6.85/SF finish (total $22–34/SF)
- Hardwood refinish: $6.75–6.85/SF
- Patch and paint (minor): $500–1,000 LS
- Painting labor: $75–85/hr
- Drywall patch (minor): $500–1,500 LS

DIV 15 — PLUMBING & HVAC:
- Install faucet (standard): $900 labor
- Install toilet (standard): $900 labor
- Install sink (standard): $900 labor
- Move plumbing lines (minor): $1,250 allowance
- Bathroom exhaust fan furnish: $225–255; install: $685
- Furnish new faucet: $350–900 (allowance)
- Furnish new toilet: $425–2,000 (allowance, varies by spec)

DIV 16 — ELECTRICAL:
- Recessed light furnish: $65–70/EA; wire & install: $325–425/EA
- Wire & install duplex outlet: $150–325/EA
- Wire & install bathroom exhaust fan: $250–275/EA
- Wire & install wall sconce: $525–575/EA
- Wire & install pendant: $625–750/EA
- Electrical permit: $400–650 LS
- Remove electrical for demo: $500–1,800 allowance
"""


# ---------------------------------------------------------------------------
# System prompt — CHG-51 + SYS-09 grounded
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = f"""You are a construction contract specialist working within a firm that follows CHG-51 (Change Event SOP) and uses SYS-09 (Firm Estimating Schedule) for pricing. Your job is to convert field notes about a scope change into a complete Potential Change Order (PCO) that a project manager can review and send to the owner for approval.

DOCUMENT TYPE:
This is a PCO — a Potential Change Order. Work described herein SHALL NOT proceed until the owner has signed and returned this document.

CAUSE CATEGORIES (use exactly one):
{chr(10).join(f'- {c}' for c in CAUSE_CATEGORIES)}

SCOPE NARRATIVE STRUCTURE (CHG-51 requirement):
1. BASELINE: What the original contract included relevant to this change.
2. CHANGE: What is different — what work is being added, modified, or deleted.
3. State precisely; write as GC scope, not as a vendor quote.

COST LINE ITEMS:
You will receive the contractor's estimated labor and material totals as anchors.
Generate a line-item cost table with these columns: Item | Cost Basis | Amount
- Break the work into logical line items by trade and task
- Reference the cost basis for each item (e.g., "SYS-09: Demo crew, 4 hrs @ $1,920/day", "SYS-09: $3.76/SF × 60 SF", "GC estimate — vendor quote")
- Line items should sum to approximately the contractor's provided totals
- If contractor totals are not provided, use SYS-09 rates to estimate
- Show SUBTOTAL, then markup lines, then TOTAL PCO AMOUNT

PRICING DATABASE — USE THESE RATES:
{SYS09_RATES}

OUTPUT FORMAT — produce these exact section headers in this order, plain text only:

SCOPE OF ADDITIONAL WORK
[Baseline paragraph, then Change paragraph]

INCLUSIONS AND EXCLUSIONS
Included: [bullet list of what is in scope]
Excluded: [bullet list of what is not in scope]

ASSUMPTIONS AND CONDITIONS
[Numbered list of access, sequencing, and pricing assumptions]

COST TABLE
[Line items in format: Description | Cost Basis | $Amount]
[End with: SUBTOTAL | | $X]
[Then: Application Fee (1%) | Per Prime Agreement | $X]
[Then: Overhead (10%) | Per Prime Agreement | $X]
[Then: Fee (10%) | Per Prime Agreement | $X]
[Then: TOTAL PCO AMOUNT | | $X]

SCHEDULE IMPACT
[Calendar days added and decision deadline tied to schedule consequence]

CAUSE OF CHANGE
[Cause category + one sentence explanation]

CONTRACT MODIFICATION
This Potential Change Order, when executed by both parties, will modify the original contract dated [original contract date] between [company name] and the Owner. All other terms and conditions of the original contract remain in full force and effect.

PROPOSED REVISED CONTRACT SUM
Original Contract Sum: $[original_contract_sum if provided, else "[To Be Confirmed]"]
This Change Order: $[TOTAL PCO AMOUNT]
Proposed Revised Contract Sum: $[sum if original provided, else "[Original Sum] + [PCO Total]"]

VALIDITY
This proposal is valid for {PCO_VALIDITY_DAYS} days from the date of issue. Pricing is subject to revision if not executed within this period.

AUTHORIZATION
No work described in this Potential Change Order shall commence until this document has been executed (signed by both parties). Verbal or written direction to proceed without a signed PCO constitutes authorization for T&M billing per the terms of the original contract.

RULES:
- Never use informal language
- Never fabricate cost numbers beyond what SYS-09 rates or contractor inputs support
- Never add scope not described in the input
- Write scope as GC scope — you control scope language and boundaries
- Output plain text only — no markdown, no # headers
- Numbers in the cost table must be internally consistent — verify your math
- This is a DRAFT for PM review"""


# ---------------------------------------------------------------------------
# Claude call
# ---------------------------------------------------------------------------

def generate_pco_text(data: dict) -> str:
    """Call Claude to generate full PCO including line-item cost table."""

    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    labor_str    = f"${data.get('labor_cost', '')}" if data.get('labor_cost') else "Not provided — use SYS-09 rates"
    material_str = f"${data.get('material_cost', '')}" if data.get('material_cost') else "Not provided — use SYS-09 rates"

    user_message = f"""Generate a PCO draft from the following field data:

Company Name: {data.get('company_name', '[TBD]')}
Project Name: {data.get('project_name', '[TBD]')}
Project Address: {data.get('project_address', '[TBD]')}
PCO Number: {data.get('change_order_number', '001')}
Date: {data.get('date', datetime.now().strftime('%B %d, %Y'))}
Original Contract Date: {data.get('original_contract_date', '[TBD]')}

Field Description of Scope Change:
{data.get('scope_description', '[No description provided]')}

Contractor's Labor Estimate (anchor for line items): {labor_str}
Contractor's Material Estimate (anchor for line items): {material_str}
Schedule Impact: {data.get('schedule_days', '[TBD]')} calendar days
Cause of Change: {data.get('reason_for_change', '[TBD]')}

Generate the complete PCO including the line-item cost table. Verify that line item amounts sum correctly to the subtotal, and that markup calculations are accurate."""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2500,
        temperature=0,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}]
    )

    return message.content[0].text


# ---------------------------------------------------------------------------
# Section parser
# ---------------------------------------------------------------------------

SECTION_HEADERS = [
    "SCOPE OF ADDITIONAL WORK",
    "INCLUSIONS AND EXCLUSIONS",
    "ASSUMPTIONS AND CONDITIONS",
    "COST TABLE",
    "SCHEDULE IMPACT",
    "CAUSE OF CHANGE",
    "CONTRACT MODIFICATION",
    "PROPOSED REVISED CONTRACT SUM",
    "VALIDITY",
    "AUTHORIZATION",
]

def parse_sections(text: str) -> dict:
    """Parse Claude's output into sections keyed by header."""
    sections = {}
    current_header = None
    current_lines = []

    for line in text.strip().split('\n'):
        stripped = line.strip().upper()
        matched = next((h for h in SECTION_HEADERS if stripped == h), None)

        if matched:
            if current_header and current_lines:
                sections[current_header] = '\n'.join(current_lines).strip()
            current_header = matched
            current_lines = []
        elif current_header:
            current_lines.append(line)

    if current_header and current_lines:
        sections[current_header] = '\n'.join(current_lines).strip()

    return sections


def parse_cost_table(cost_text: str) -> list:
    """Parse the cost table section into a list of [description, basis, amount] rows."""
    rows = []
    for line in cost_text.strip().split('\n'):
        line = line.strip()
        if not line:
            continue
        parts = [p.strip() for p in line.split('|')]
        if len(parts) >= 3:
            rows.append(parts[:3])
        elif len(parts) == 2:
            rows.append([parts[0], '', parts[1]])
        elif len(parts) == 1 and parts[0]:
            rows.append([parts[0], '', ''])
    return rows


# ---------------------------------------------------------------------------
# PDF builder
# ---------------------------------------------------------------------------

def build_pdf(data: dict, sections: dict, output_path: str):
    """Build professional PDF PCO with line-item cost table."""

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.875 * inch,
        leftMargin=0.875 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch
    )

    company_style = ParagraphStyle(
        'Company', fontSize=14, fontName='Helvetica-Bold',
        alignment=TA_LEFT, spaceAfter=4
    )
    title_style = ParagraphStyle(
        'Title', fontSize=18, fontName='Helvetica-Bold',
        alignment=TA_CENTER, spaceBefore=8, spaceAfter=2
    )
    subtitle_style = ParagraphStyle(
        'Subtitle', fontSize=10, fontName='Helvetica',
        alignment=TA_CENTER, spaceAfter=4,
        textColor=colors.HexColor('#555555')
    )
    draft_style = ParagraphStyle(
        'Draft', fontSize=9, fontName='Helvetica-Bold',
        alignment=TA_CENTER, textColor=colors.red, spaceAfter=12
    )
    section_header_style = ParagraphStyle(
        'SectionHeader', fontSize=9, fontName='Helvetica-Bold',
        spaceBefore=12, spaceAfter=3,
        textColor=colors.HexColor('#333333')
    )
    body_style = ParagraphStyle(
        'Body', fontSize=9, fontName='Helvetica',
        spaceAfter=4, leading=13
    )

    story = []

    # Letterhead
    story.append(Paragraph(data.get('company_name', ''), company_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.black, spaceAfter=4))
    story.append(Paragraph("POTENTIAL CHANGE ORDER", title_style))
    story.append(Paragraph("For Owner Review and Approval", subtitle_style))
    story.append(Paragraph("DRAFT — FOR PM REVIEW ONLY", draft_style))

    # Header info
    co_date      = data.get('date', datetime.now().strftime('%B %d, %Y'))
    validity_date = (datetime.now() + timedelta(days=PCO_VALIDITY_DAYS)).strftime('%B %d, %Y')

    header_rows = [
        [f"PCO No.: {data.get('change_order_number', '001')}",
         f"Date: {co_date}"],
        [f"Project: {data.get('project_name', '')}",
         f"Contract Date: {data.get('original_contract_date', '[TBD]')}"],
        [f"Address: {data.get('project_address', '')}",
         f"Valid Through: {validity_date}"],
    ]
    header_table = Table(header_rows, colWidths=[3.5 * inch, 3.25 * inch])
    header_table.setStyle(TableStyle([
        ('FONTNAME',    (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE',    (0, 0), (-1, -1), 9),
        ('VALIGN',      (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING',  (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey, spaceAfter=2))

    # Sections — special handling for COST TABLE
    for header in SECTION_HEADERS:
        if header not in sections:
            continue

        story.append(Paragraph(header, section_header_style))

        if header == "COST TABLE":
            rows = parse_cost_table(sections[header])
            if rows:
                table_data = [["DESCRIPTION", "COST BASIS", "AMOUNT"]]
                for row in rows:
                    desc   = row[0] if len(row) > 0 else ''
                    basis  = row[1] if len(row) > 1 else ''
                    amount = row[2] if len(row) > 2 else ''
                    table_data.append([desc, basis, amount])

                col_widths = [2.8 * inch, 2.5 * inch, 1.45 * inch]
                cost_table = Table(table_data, colWidths=col_widths)
                cost_table.setStyle(TableStyle([
                    ('FONTNAME',     (0, 0),  (-1, 0),  'Helvetica-Bold'),
                    ('FONTNAME',     (0, 1),  (-1, -1), 'Helvetica'),
                    ('FONTSIZE',     (0, 0),  (-1, -1), 8),
                    ('BACKGROUND',   (0, 0),  (-1, 0),  colors.HexColor('#1a1a18')),
                    ('TEXTCOLOR',    (0, 0),  (-1, 0),  colors.white),
                    ('ALIGN',        (2, 0),  (2, -1),  'RIGHT'),
                    ('TOPPADDING',   (0, 0),  (-1, -1), 3),
                    ('BOTTOMPADDING',(0, 0),  (-1, -1), 3),
                    ('ROWBACKGROUNDS',(0, 1), (-1, -1),
                     [colors.HexColor('#f4f2ee'), colors.white]),
                    ('LINEBELOW',    (0, 0),  (-1, 0),  1, colors.black),
                    ('BOX',          (0, 0),  (-1, -1), 0.5, colors.HexColor('#d4cfc8')),
                ]))
                story.append(cost_table)
        else:
            content = sections[header].replace('\n', '<br/>')
            story.append(Paragraph(content, body_style))

    # Signature blocks
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey, spaceAfter=8))

    sig_rows = [
        ["Contractor:", "_" * 34, "Date:", "_" * 14],
        ["", "", "", ""],
        ["Print Name:", "_" * 31, "", ""],
        ["", "", "", ""],
        ["Owner:", "_" * 37, "Date:", "_" * 14],
        ["", "", "", ""],
        ["Print Name:", "_" * 31, "", ""],
    ]
    sig_table = Table(sig_rows, colWidths=[1.0 * inch, 2.4 * inch, 0.6 * inch, 2.75 * inch])
    sig_table.setStyle(TableStyle([
        ('FONTNAME',      (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE',      (0, 0), (-1, -1), 9),
        ('TOPPADDING',    (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    story.append(sig_table)

    doc.build(story)


# ---------------------------------------------------------------------------
# Word doc builder
# ---------------------------------------------------------------------------

def _add_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_word(data: dict, sections: dict, output_path: str):
    """Build editable Word doc PCO."""

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Company
    p = doc.add_paragraph()
    run = p.add_run(data.get('company_name', ''))
    run.bold = True
    run.font.size = Pt(14)
    _add_rule(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("POTENTIAL CHANGE ORDER")
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("For Owner Review and Approval")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DRAFT — FOR PM REVIEW ONLY")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.add_paragraph()

    co_date       = data.get('date', datetime.now().strftime('%B %d, %Y'))
    validity_date = (datetime.now() + timedelta(days=PCO_VALIDITY_DAYS)).strftime('%B %d, %Y')

    for line in [
        f"PCO No.: {data.get('change_order_number', '001')}",
        f"Date: {co_date}",
        f"Project: {data.get('project_name', '')}",
        f"Address: {data.get('project_address', '')}",
        f"Original Contract Date: {data.get('original_contract_date', '[TBD]')}",
        f"Valid Through: {validity_date}",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)

    doc.add_paragraph()
    _add_rule(doc)

    for header in SECTION_HEADERS:
        if header not in sections:
            continue

        p = doc.add_paragraph()
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)

        if header == "COST TABLE":
            rows = parse_cost_table(sections[header])
            if rows:
                table = doc.add_table(rows=len(rows) + 1, cols=3)
                table.style = 'Table Grid'

                hdr = table.rows[0]
                for i, text in enumerate(["DESCRIPTION", "COST BASIS", "AMOUNT"]):
                    hdr.cells[i].text = text
                    for run in hdr.cells[i].paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)

                for i, row_data in enumerate(rows):
                    row = table.rows[i + 1]
                    for j, val in enumerate(row_data[:3]):
                        row.cells[j].text = val
                        for run in row.cells[j].paragraphs[0].runs:
                            run.font.size = Pt(9)
        else:
            p = doc.add_paragraph()
            run = p.add_run(sections[header])
            run.font.size = Pt(10)

        doc.add_paragraph()

    # Signatures
    doc.add_paragraph()
    for line in [
        "Contractor: ____________________________________________   Date: ________________",
        "",
        "Print Name: ___________________________________________",
        "",
        "",
        "Owner: _________________________________________________   Date: ________________",
        "",
        "Print Name: ___________________________________________",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate(data: dict, output_dir: str = "./output") -> dict:
    """
    Generate a PCO from structured field data.
    Claude generates the full document including line-item cost table.

    Args:
        data: dict with company_name, project_name, project_address,
              change_order_number, date, original_contract_date,
              scope_description, labor_cost, material_cost,
              schedule_days, reason_for_change

    Returns:
        dict with pdf, word, generated_text, sections
    """

    Path(output_dir).mkdir(parents=True, exist_ok=True)

    co_num   = data.get('change_order_number', '001')
    slug     = data.get('project_name', 'project').replace(' ', '_').lower()[:30]
    date_str = datetime.now().strftime('%Y%m%d')
    base     = f"pco_{co_num}_{slug}_{date_str}"

    pdf_path  = os.path.join(output_dir, f"{base}.pdf")
    word_path = os.path.join(output_dir, f"{base}.docx")

    print("Calling Claude (CHG-51 + SYS-09 grounded)...")
    generated_text = generate_pco_text(data)

    print("Parsing sections...")
    sections = parse_sections(generated_text)

    print(f"Building PDF  → {pdf_path}")
    build_pdf(data, sections, pdf_path)

    print(f"Building Word → {word_path}")
    build_word(data, sections, word_path)

    print("Done.")
    return {
        "pdf":            pdf_path,
        "word":           word_path,
        "generated_text": generated_text,
        "sections":       sections,
    }
